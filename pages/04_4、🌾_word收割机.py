import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re  # 正则支持

# --- 页面配置 ---
st.set_page_config(page_title="Word收割机", page_icon="🌾", layout="wide")

st.title("🌾 Word收割机 (全学历覆盖版)")
st.markdown("""
### 工业级批量处理方案
- ✅ **全学历支持**：自动拆分 本科 / 硕士 / 博士
- ✅ **批量稳健性**：支持 60+ 文件批量上传，单文件报错不影响整体
- ✅ **智能去噪**：自动过滤干扰行、修复合并单元格误判
""")

# ==========================================
# 🔧 核心解析逻辑
# ==========================================

def get_cell_text_right_of(table, label_keywords):
    """
    在表格中查找包含指定关键词的单元格，并返回它【右边】那个单元格的内容。
    (包含防误判、防合并、防同单元格挤压的逻辑)
    """
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell_text = cell.text.strip().replace(" ", "") 
            cell_text_raw = cell.text.strip()
            
            # 检查关键词
            matched_keyword = next((k for k in label_keywords if k in cell_text), None)
            
            if matched_keyword:
                # --- 情况 A: 标签在左，值在右 ---
                if i + 1 < len(row.cells):
                    next_cell = row.cells[i+1]
                    val = next_cell.text.strip()
                    
                    # 🛑 防御逻辑 1: 左右是同一个格子 (合并单元格) -> 跳过
                    if cell is next_cell:
                        pass 
                    
                    # 🛑 防御逻辑 2: 值就是标签本身 (表头占位) -> 跳过
                    elif any(k in val.replace(" ", "") for k in label_keywords):
                         pass
                    
                    # ✅ 提取有效值
                    elif val: 
                        return val

                # --- 情况 B: 标签和值在同一个格子里 (正则分割) ---
                if len(cell_text) > len(matched_keyword) + 1:
                    parts = re.split(r'[:：]', cell_text_raw)
                    if len(parts) > 1:
                        return parts[1].strip()

    return "" 

def extract_study_resume_structured(doc):
    """
    查找“学习简历”表格，精准提取 本科/硕士/博士 信息。
    核心升级：增加博士字段支持。
    """
    # 预置空字段 (确保所有列都存在，方便Excel对齐)
    resume_info = {
        "博士-时间": "", "博士-院校": "", "博士-专业": "",
        "硕士-时间": "", "硕士-院校": "", "硕士-专业": "",
        "本科-时间": "", "本科-院校": "", "本科-专业": ""
    }
    
    for table in doc.tables:
        if len(table.rows) > 0:
            header_text = "".join([cell.text for cell in table.rows[0].cells])
            
            if "起止年月" in header_text and "毕业院校" in header_text:
                for row in table.rows[1:]:
                    # 🛑 排雷：跳过“填写说明”行
                    row_full_text = "".join([c.text for c in row.cells])
                    if "填写说明" in row_full_text:
                        continue 
                    
                    cells = row.cells
                    if len(cells) >= 5:
                        date = cells[0].text.strip()    # 第1列: 时间
                        school = cells[1].text.strip()  # 第2列: 学校
                        major = cells[2].text.strip()   # 第3列: 专业
                        degree = cells[4].text.strip()  # 第5列: 学位
                        
                        if degree:
                            # --- 🎓 三级分流逻辑 ---
                            if "博士" in degree:
                                resume_info["博士-时间"] = date
                                resume_info["博士-院校"] = school
                                resume_info["博士-专业"] = major

                            elif "硕士" in degree:
                                resume_info["硕士-时间"] = date
                                resume_info["硕士-院校"] = school
                                resume_info["硕士-专业"] = major

                            elif "学士" in degree or "本科" in degree:
                                resume_info["本科-时间"] = date
                                resume_info["本科-院校"] = school
                                resume_info["本科-专业"] = major
                            
                break 
    
    return resume_info

def extract_info_from_docx(file):
    """
    主解析函数
    """
    doc = Document(file)
    info = {}
    
    # 1. 基本信息
    if len(doc.tables) > 0:
        base_table = doc.tables[0]
        info['姓名'] = get_cell_text_right_of(base_table, ['姓名'])
        info['性别'] = get_cell_text_right_of(base_table, ['性别'])
        info['出生日期'] = get_cell_text_right_of(base_table, ['出生日期'])
        info['民族'] = get_cell_text_right_of(base_table, ['民族'])
        info['出生地'] = get_cell_text_right_of(base_table, ['出生地'])
        info['手机号'] = get_cell_text_right_of(base_table, ['手机', '联系电话'])
        info['邮箱'] = get_cell_text_right_of(base_table, ['邮箱', 'E-mail'])
        info['入职时间'] = get_cell_text_right_of(base_table, ['入职时间'])
    else:
        info['状态'] = "未找到表格"

    # 2. 学历信息 (含博士)
    resume_data = extract_study_resume_structured(doc)
    info.update(resume_data)

    return info

# ==========================================
# 🚀 Streamlit 主程序
# ==========================================

uploaded_files = st.file_uploader(
    "📂 请上传员工信息登记表 (支持多选，建议单次不超过100个)", 
    type=['docx'], 
    accept_multiple_files=True
)

if uploaded_files:
    total_files = len(uploaded_files)
    st.info(f"⚡ 已接收 {total_files} 个文件，引擎全速运转中...")
    
    all_data = []
    failed_files = [] # 记录失败的文件
    
    # 进度条
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # --- 批量循环 ---
    for i, file in enumerate(uploaded_files):
        try:
            # 实时更新状态文字
            status_text.text(f"正在处理 ({i+1}/{total_files}): {file.name}")
            
            single_data = extract_info_from_docx(file)
            single_data['源文件名'] = file.name 
            all_data.append(single_data)
            
        except Exception as e:
            # 🛡️ 错误隔离
            failed_files.append(f"{file.name}: {str(e)}")
            all_data.append({'源文件名': file.name, '姓名': '解析失败'})
        
        # 更新进度
        progress_bar.progress((i + 1) / total_files)
    
    status_text.text("✅ 所有文件处理完毕！")
    
    # --- 结果处理 ---
    df = pd.DataFrame(all_data)
    
    # 🌟 关键修改：让索引从1开始
    df.index = df.index + 1
    
    # 优化列排序 (博士优先展示)
    priority_cols = [
        '姓名', '性别', '民族', '手机号', '邮箱',
        '博士-院校', '博士-专业', '博士-时间',
        '硕士-院校', '硕士-专业', '硕士-时间', 
        '本科-院校', '本科-专业', '本科-时间'
    ]
    other_cols = [c for c in df.columns if c not in priority_cols and c != '源文件名']
    final_cols = [c for c in (priority_cols + other_cols) if c in df.columns]
    
    df = df[final_cols]

    st.markdown("---")
    
    # 📊 如果有失败的文件，展示错误报告
    if failed_files:
        with st.expander(f"⚠️ 注意：有 {len(failed_files)} 个文件处理出现异常", expanded=True):
            for err in failed_files:
                st.write(err)

    st.subheader(f"📊 汇总结果 ({len(all_data)}条数据)")
    
    # 自由选择列
    all_columns = df.columns.tolist()
    selected_cols = st.multiselect(
        "✨ 自定义导出字段:",
        options=all_columns,
        default=all_columns 
    )
    
    if selected_cols:
        result_df = df[selected_cols]
        st.dataframe(result_df)
        
        # 导出 Excel
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            result_df.to_excel(writer, index=False, sheet_name='汇总数据')
        
        st.download_button(
            label="📥 下载 Excel 汇总表",
            data=output.getvalue(),
            file_name="员工信息全量汇总(含博士).xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        st.warning("请至少勾选一个字段。")