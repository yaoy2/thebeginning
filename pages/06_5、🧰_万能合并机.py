import streamlit as st
import pandas as pd
from docx import Document
from pypdf import PdfReader
from io import BytesIO
from utils.ui_theme import render_home_link

# --- 页面配置 ---
st.set_page_config(page_title="万能合并机", page_icon="🧰", layout="wide")
render_home_link()

st.title("🧰 万能文档/表格合并机 (V2.0 破甲版)")
st.markdown("""
### 💡 True工具箱
你只管扔，我全包：
1. **Word/PDF** ➡️ 自动提取文字，合并成 **1个 Word 文档** (支持解除 PDF 复制限制)。
2. **Excel** ➡️ 自动搬运，合并成 **1个 Excel 的多个 Sheet** (数据归档神器)。
""")

# ==========================================
# 🔧 引擎 1: 文档合并 (Word + PDF -> Word)
# ==========================================
def merge_to_docx(files):
    merged_doc = Document()
    
    # 进度条
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_files = len(files)
    
    for i, file in enumerate(files):
        status_text.text(f"正在提取 ({i+1}/{total_files}): {file.name}")
        
        # 插入一级标题 (来源) - 方便 Gem 识别出处
        merged_doc.add_heading(f"📄 文件来源：{file.name}", level=1)
        
        file_ext = file.name.split('.')[-1].lower()
        
        try:
            # --- A. 处理 Word ---
            if file_ext == 'docx':
                sub_doc = Document(file)
                for para in sub_doc.paragraphs:
                    if para.text.strip():
                        merged_doc.add_paragraph(para.text)
            
            # --- B. 处理 PDF (含破甲逻辑) ---
            elif file_ext == 'pdf':
                reader = PdfReader(file)
                
                # 🛑 破甲逻辑：如果文件被权限加密 (能看不能复制)
                if reader.is_encrypted:
                    try:
                        # 尝试用空密码解密 (绝大多数行政文件的加密方式)
                        reader.decrypt('')
                    except:
                        merged_doc.add_paragraph(f"[⚠️ 警告：文件 {file.name} 有强密码保护，无法提取]")
                        continue

                # 遍历每一页提取文字
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        merged_doc.add_paragraph(text)
                    else:
                        # 如果提取出来是空的，可能是纯图片扫描件
                        merged_doc.add_paragraph("[⚠️ 提示：此页可能是纯图片扫描件，未提取到文字]")
                        
        except Exception as e:
            merged_doc.add_paragraph(f"[读取失败: {str(e)}]")
            st.warning(f"文件 {file.name} 读取出现问题: {e}")

        # 插入分页符，让不同文件的内容在视觉上隔开
        merged_doc.add_page_break()
        progress_bar.progress((i + 1) / total_files)
        
    status_text.text("✅ 文档合并完成！")
    return merged_doc

# ==========================================
# 🔧 引擎 2: 表格合并 (Excels -> Sheets)
# ==========================================
def merge_to_excel_sheets(files):
    output = BytesIO()
    
    # 使用 xlsxwriter 引擎
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        total_files = len(files)
        
        # 记录用过的 sheet 名，防止重复
        used_sheet_names = {}
        
        for i, file in enumerate(files):
            status_text.text(f"正在搬运 ({i+1}/{total_files}): {file.name}")
            
            try:
                # 读取 Excel
                df = pd.read_excel(file)
                
                # --- Sheet 名处理逻辑 ---
                # Excel sheet名最长31字符
                base_name = file.name.split('.')[0][:25] 
                # 去除非法字符
                invalid_chars = ['[', ']', ':', '*', '?', '/', '\\']
                for char in invalid_chars:
                    base_name = base_name.replace(char, '_')
                
                # 处理重名 (如：名单.xlsx, 名单.xlsx -> 名单, 名单_1)
                if base_name in used_sheet_names:
                    used_sheet_names[base_name] += 1
                    sheet_name = f"{base_name}_{used_sheet_names[base_name]}"
                else:
                    used_sheet_names[base_name] = 0
                    sheet_name = base_name
                
                # 写入新的 Sheet
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                
            except Exception as e:
                st.error(f"表格 {file.name} 处理失败: {e}")
            
            progress_bar.progress((i + 1) / total_files)
            
    status_text.text("✅ 表格合并完成！")
    return output

# ==========================================
# 🚀 主控台布局
# ==========================================

st.markdown("---")

col1, col2 = st.columns(2)

# --- 左侧：文档区 ---
with col1:
    st.subheader("📄 文档处理区")
    st.caption("支持 .docx, .pdf (含加密) | 产出：汇总版 Word")
    doc_files = st.file_uploader("拖入文档", type=['docx', 'pdf'], accept_multiple_files=True, key="doc_uploader")
    
    if doc_files:
        if st.button("开始合并文档", type="primary"):
            final_doc = merge_to_docx(doc_files)
            
            # 保存
            doc_buffer = BytesIO()
            final_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.download_button(
                "📥 下载超级文档 (.docx)",
                data=doc_buffer,
                file_name="全院资料汇总_AI版.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# --- 右侧：表格区 ---
with col2:
    st.subheader("📊 表格处理区")
    st.caption("支持 .xlsx | 产出：多 Sheet 版 Excel")
    excel_files = st.file_uploader("拖入表格", type=['xlsx', 'xls'], accept_multiple_files=True, key="xls_uploader")
    
    if excel_files:
        if st.button("开始合并表格"):
            final_excel = merge_to_excel_sheets(excel_files)
            
            st.download_button(
                "📥 下载超级表格 (.xlsx)",
                data=final_excel.getvalue(),
                file_name="全院表格汇总_多Sheet版.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
