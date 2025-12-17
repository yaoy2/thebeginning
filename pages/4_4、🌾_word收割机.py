import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re  # æ­£åˆ™æ”¯æŒ

# --- é¡µé¢é…ç½® ---
st.set_page_config(page_title="Wordæ”¶å‰²æœº", page_icon="ğŸŒ¾", layout="wide")

st.title("ğŸŒ¾ Wordæ”¶å‰²æœº (å…¨å­¦å†è¦†ç›–ç‰ˆ)")
st.markdown("""
### å·¥ä¸šçº§æ‰¹é‡å¤„ç†æ–¹æ¡ˆ
- âœ… **å…¨å­¦å†æ”¯æŒ**ï¼šè‡ªåŠ¨æ‹†åˆ† æœ¬ç§‘ / ç¡•å£« / åšå£«
- âœ… **æ‰¹é‡ç¨³å¥æ€§**ï¼šæ”¯æŒ 60+ æ–‡ä»¶æ‰¹é‡ä¸Šä¼ ï¼Œå•æ–‡ä»¶æŠ¥é”™ä¸å½±å“æ•´ä½“
- âœ… **æ™ºèƒ½å»å™ª**ï¼šè‡ªåŠ¨è¿‡æ»¤å¹²æ‰°è¡Œã€ä¿®å¤åˆå¹¶å•å…ƒæ ¼è¯¯åˆ¤
""")

# ==========================================
# ğŸ”§ æ ¸å¿ƒè§£æé€»è¾‘
# ==========================================

def get_cell_text_right_of(table, label_keywords):
    """
    åœ¨è¡¨æ ¼ä¸­æŸ¥æ‰¾åŒ…å«æŒ‡å®šå…³é”®è¯çš„å•å…ƒæ ¼ï¼Œå¹¶è¿”å›å®ƒã€å³è¾¹ã€‘é‚£ä¸ªå•å…ƒæ ¼çš„å†…å®¹ã€‚
    (åŒ…å«é˜²è¯¯åˆ¤ã€é˜²åˆå¹¶ã€é˜²åŒå•å…ƒæ ¼æŒ¤å‹çš„é€»è¾‘)
    """
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell_text = cell.text.strip().replace(" ", "") 
            cell_text_raw = cell.text.strip()
            
            # æ£€æŸ¥å…³é”®è¯
            matched_keyword = next((k for k in label_keywords if k in cell_text), None)
            
            if matched_keyword:
                # --- æƒ…å†µ A: æ ‡ç­¾åœ¨å·¦ï¼Œå€¼åœ¨å³ ---
                if i + 1 < len(row.cells):
                    next_cell = row.cells[i+1]
                    val = next_cell.text.strip()
                    
                    # ğŸ›‘ é˜²å¾¡é€»è¾‘ 1: å·¦å³æ˜¯åŒä¸€ä¸ªæ ¼å­ (åˆå¹¶å•å…ƒæ ¼) -> è·³è¿‡
                    if cell is next_cell:
                        pass 
                    
                    # ğŸ›‘ é˜²å¾¡é€»è¾‘ 2: å€¼å°±æ˜¯æ ‡ç­¾æœ¬èº« (è¡¨å¤´å ä½) -> è·³è¿‡
                    elif any(k in val.replace(" ", "") for k in label_keywords):
                         pass
                    
                    # âœ… æå–æœ‰æ•ˆå€¼
                    elif val: 
                        return val

                # --- æƒ…å†µ B: æ ‡ç­¾å’Œå€¼åœ¨åŒä¸€ä¸ªæ ¼å­é‡Œ (æ­£åˆ™åˆ†å‰²) ---
                if len(cell_text) > len(matched_keyword) + 1:
                    parts = re.split(r'[:ï¼š]', cell_text_raw)
                    if len(parts) > 1:
                        return parts[1].strip()

    return "" 

def extract_study_resume_structured(doc):
    """
    æŸ¥æ‰¾â€œå­¦ä¹ ç®€å†â€è¡¨æ ¼ï¼Œç²¾å‡†æå– æœ¬ç§‘/ç¡•å£«/åšå£« ä¿¡æ¯ã€‚
    æ ¸å¿ƒå‡çº§ï¼šå¢åŠ åšå£«å­—æ®µæ”¯æŒã€‚
    """
    # é¢„ç½®ç©ºå­—æ®µ (ç¡®ä¿æ‰€æœ‰åˆ—éƒ½å­˜åœ¨ï¼Œæ–¹ä¾¿Excelå¯¹é½)
    resume_info = {
        "åšå£«-æ—¶é—´": "", "åšå£«-é™¢æ ¡": "", "åšå£«-ä¸“ä¸š": "",
        "ç¡•å£«-æ—¶é—´": "", "ç¡•å£«-é™¢æ ¡": "", "ç¡•å£«-ä¸“ä¸š": "",
        "æœ¬ç§‘-æ—¶é—´": "", "æœ¬ç§‘-é™¢æ ¡": "", "æœ¬ç§‘-ä¸“ä¸š": ""
    }
    
    for table in doc.tables:
        if len(table.rows) > 0:
            header_text = "".join([cell.text for cell in table.rows[0].cells])
            
            if "èµ·æ­¢å¹´æœˆ" in header_text and "æ¯•ä¸šé™¢æ ¡" in header_text:
                for row in table.rows[1:]:
                    # ğŸ›‘ æ’é›·ï¼šè·³è¿‡â€œå¡«å†™è¯´æ˜â€è¡Œ
                    row_full_text = "".join([c.text for c in row.cells])
                    if "å¡«å†™è¯´æ˜" in row_full_text:
                        continue 
                    
                    cells = row.cells
                    if len(cells) >= 5:
                        date = cells[0].text.strip()    # ç¬¬1åˆ—: æ—¶é—´
                        school = cells[1].text.strip()  # ç¬¬2åˆ—: å­¦æ ¡
                        major = cells[2].text.strip()   # ç¬¬3åˆ—: ä¸“ä¸š
                        degree = cells[4].text.strip()  # ç¬¬5åˆ—: å­¦ä½
                        
                        if degree:
                            # --- ğŸ“ ä¸‰çº§åˆ†æµé€»è¾‘ ---
                            if "åšå£«" in degree:
                                resume_info["åšå£«-æ—¶é—´"] = date
                                resume_info["åšå£«-é™¢æ ¡"] = school
                                resume_info["åšå£«-ä¸“ä¸š"] = major

                            elif "ç¡•å£«" in degree:
                                resume_info["ç¡•å£«-æ—¶é—´"] = date
                                resume_info["ç¡•å£«-é™¢æ ¡"] = school
                                resume_info["ç¡•å£«-ä¸“ä¸š"] = major

                            elif "å­¦å£«" in degree or "æœ¬ç§‘" in degree:
                                resume_info["æœ¬ç§‘-æ—¶é—´"] = date
                                resume_info["æœ¬ç§‘-é™¢æ ¡"] = school
                                resume_info["æœ¬ç§‘-ä¸“ä¸š"] = major
                            
                break 
    
    return resume_info

def extract_info_from_docx(file):
    """
    ä¸»è§£æå‡½æ•°
    """
    doc = Document(file)
    info = {}
    
    # 1. åŸºæœ¬ä¿¡æ¯
    if len(doc.tables) > 0:
        base_table = doc.tables[0]
        info['å§“å'] = get_cell_text_right_of(base_table, ['å§“å'])
        info['æ€§åˆ«'] = get_cell_text_right_of(base_table, ['æ€§åˆ«'])
        info['å‡ºç”Ÿæ—¥æœŸ'] = get_cell_text_right_of(base_table, ['å‡ºç”Ÿæ—¥æœŸ'])
        info['æ°‘æ—'] = get_cell_text_right_of(base_table, ['æ°‘æ—'])
        info['å‡ºç”Ÿåœ°'] = get_cell_text_right_of(base_table, ['å‡ºç”Ÿåœ°'])
        info['æ‰‹æœºå·'] = get_cell_text_right_of(base_table, ['æ‰‹æœº', 'è”ç³»ç”µè¯'])
        info['é‚®ç®±'] = get_cell_text_right_of(base_table, ['é‚®ç®±', 'E-mail'])
        info['å…¥èŒæ—¶é—´'] = get_cell_text_right_of(base_table, ['å…¥èŒæ—¶é—´'])
    else:
        info['çŠ¶æ€'] = "æœªæ‰¾åˆ°è¡¨æ ¼"

    # 2. å­¦å†ä¿¡æ¯ (å«åšå£«)
    resume_data = extract_study_resume_structured(doc)
    info.update(resume_data)

    return info

# ==========================================
# ğŸš€ Streamlit ä¸»ç¨‹åº
# ==========================================

uploaded_files = st.file_uploader(
    "ğŸ“‚ è¯·ä¸Šä¼ å‘˜å·¥ä¿¡æ¯ç™»è®°è¡¨ (æ”¯æŒå¤šé€‰ï¼Œå»ºè®®å•æ¬¡ä¸è¶…è¿‡100ä¸ª)", 
    type=['docx'], 
    accept_multiple_files=True
)

if uploaded_files:
    total_files = len(uploaded_files)
    st.info(f"âš¡ å·²æ¥æ”¶ {total_files} ä¸ªæ–‡ä»¶ï¼Œå¼•æ“å…¨é€Ÿè¿è½¬ä¸­...")
    
    all_data = []
    failed_files = [] # è®°å½•å¤±è´¥çš„æ–‡ä»¶
    
    # è¿›åº¦æ¡
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # --- æ‰¹é‡å¾ªç¯ ---
    for i, file in enumerate(uploaded_files):
        try:
            # å®æ—¶æ›´æ–°çŠ¶æ€æ–‡å­—
            status_text.text(f"æ­£åœ¨å¤„ç† ({i+1}/{total_files}): {file.name}")
            
            single_data = extract_info_from_docx(file)
            single_data['æºæ–‡ä»¶å'] = file.name 
            all_data.append(single_data)
            
        except Exception as e:
            # ğŸ›¡ï¸ é”™è¯¯éš”ç¦»
            failed_files.append(f"{file.name}: {str(e)}")
            all_data.append({'æºæ–‡ä»¶å': file.name, 'å§“å': 'è§£æå¤±è´¥'})
        
        # æ›´æ–°è¿›åº¦
        progress_bar.progress((i + 1) / total_files)
    
    status_text.text("âœ… æ‰€æœ‰æ–‡ä»¶å¤„ç†å®Œæ¯•ï¼")
    
    # --- ç»“æœå¤„ç† ---
    df = pd.DataFrame(all_data)
    
    # ğŸŒŸ å…³é”®ä¿®æ”¹ï¼šè®©ç´¢å¼•ä»1å¼€å§‹
    df.index = df.index + 1
    
    # ä¼˜åŒ–åˆ—æ’åº (åšå£«ä¼˜å…ˆå±•ç¤º)
    priority_cols = [
        'å§“å', 'æ€§åˆ«', 'æ°‘æ—', 'æ‰‹æœºå·', 'é‚®ç®±',
        'åšå£«-é™¢æ ¡', 'åšå£«-ä¸“ä¸š', 'åšå£«-æ—¶é—´',
        'ç¡•å£«-é™¢æ ¡', 'ç¡•å£«-ä¸“ä¸š', 'ç¡•å£«-æ—¶é—´', 
        'æœ¬ç§‘-é™¢æ ¡', 'æœ¬ç§‘-ä¸“ä¸š', 'æœ¬ç§‘-æ—¶é—´'
    ]
    other_cols = [c for c in df.columns if c not in priority_cols and c != 'æºæ–‡ä»¶å']
    final_cols = [c for c in (priority_cols + other_cols) if c in df.columns]
    
    df = df[final_cols]

    st.markdown("---")
    
    # ğŸ“Š å¦‚æœæœ‰å¤±è´¥çš„æ–‡ä»¶ï¼Œå±•ç¤ºé”™è¯¯æŠ¥å‘Š
    if failed_files:
        with st.expander(f"âš ï¸ æ³¨æ„ï¼šæœ‰ {len(failed_files)} ä¸ªæ–‡ä»¶å¤„ç†å‡ºç°å¼‚å¸¸", expanded=True):
            for err in failed_files:
                st.write(err)

    st.subheader(f"ğŸ“Š æ±‡æ€»ç»“æœ ({len(all_data)}æ¡æ•°æ®)")
    
    # è‡ªç”±é€‰æ‹©åˆ—
    all_columns = df.columns.tolist()
    selected_cols = st.multiselect(
        "âœ¨ è‡ªå®šä¹‰å¯¼å‡ºå­—æ®µ:",
        options=all_columns,
        default=all_columns 
    )
    
    if selected_cols:
        result_df = df[selected_cols]
        st.dataframe(result_df)
        
        # å¯¼å‡º Excel
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            result_df.to_excel(writer, index=False, sheet_name='æ±‡æ€»æ•°æ®')
        
        st.download_button(
            label="ğŸ“¥ ä¸‹è½½ Excel æ±‡æ€»è¡¨",
            data=output.getvalue(),
            file_name="å‘˜å·¥ä¿¡æ¯å…¨é‡æ±‡æ€»(å«åšå£«).xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        st.warning("è¯·è‡³å°‘å‹¾é€‰ä¸€ä¸ªå­—æ®µã€‚")