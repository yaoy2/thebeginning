import streamlit as st
import pandas as pd
from docx import Document
from pypdf import PdfReader
from io import BytesIO

# --- é¡µé¢é…ç½® ---
st.set_page_config(page_title="ä¸‡èƒ½åˆå¹¶æœº", page_icon="ğŸ§°", layout="wide")

st.title("ğŸ§° ä¸‡èƒ½æ–‡æ¡£/è¡¨æ ¼åˆå¹¶æœº (V2.0 ç ´ç”²ç‰ˆ)")
st.markdown("""
### ğŸ’¡ Trueå·¥å…·ç®±
ä½ åªç®¡æ‰”ï¼Œæˆ‘å…¨åŒ…ï¼š
1. **Word/PDF** â¡ï¸ è‡ªåŠ¨æå–æ–‡å­—ï¼Œåˆå¹¶æˆ **1ä¸ª Word æ–‡æ¡£** (æ”¯æŒè§£é™¤ PDF å¤åˆ¶é™åˆ¶)ã€‚
2. **Excel** â¡ï¸ è‡ªåŠ¨æ¬è¿ï¼Œåˆå¹¶æˆ **1ä¸ª Excel çš„å¤šä¸ª Sheet** (æ•°æ®å½’æ¡£ç¥å™¨)ã€‚
""")

# ==========================================
# ğŸ”§ å¼•æ“ 1: æ–‡æ¡£åˆå¹¶ (Word + PDF -> Word)
# ==========================================
def merge_to_docx(files):
    merged_doc = Document()
    
    # è¿›åº¦æ¡
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_files = len(files)
    
    for i, file in enumerate(files):
        status_text.text(f"æ­£åœ¨æå– ({i+1}/{total_files}): {file.name}")
        
        # æ’å…¥ä¸€çº§æ ‡é¢˜ (æ¥æº) - æ–¹ä¾¿ Gem è¯†åˆ«å‡ºå¤„
        merged_doc.add_heading(f"ğŸ“„ æ–‡ä»¶æ¥æºï¼š{file.name}", level=1)
        
        file_ext = file.name.split('.')[-1].lower()
        
        try:
            # --- A. å¤„ç† Word ---
            if file_ext == 'docx':
                sub_doc = Document(file)
                for para in sub_doc.paragraphs:
                    if para.text.strip():
                        merged_doc.add_paragraph(para.text)
            
            # --- B. å¤„ç† PDF (å«ç ´ç”²é€»è¾‘) ---
            elif file_ext == 'pdf':
                reader = PdfReader(file)
                
                # ğŸ›‘ ç ´ç”²é€»è¾‘ï¼šå¦‚æœæ–‡ä»¶è¢«æƒé™åŠ å¯† (èƒ½çœ‹ä¸èƒ½å¤åˆ¶)
                if reader.is_encrypted:
                    try:
                        # å°è¯•ç”¨ç©ºå¯†ç è§£å¯† (ç»å¤§å¤šæ•°è¡Œæ”¿æ–‡ä»¶çš„åŠ å¯†æ–¹å¼)
                        reader.decrypt('')
                    except:
                        merged_doc.add_paragraph(f"[âš ï¸ è­¦å‘Šï¼šæ–‡ä»¶ {file.name} æœ‰å¼ºå¯†ç ä¿æŠ¤ï¼Œæ— æ³•æå–]")
                        continue

                # éå†æ¯ä¸€é¡µæå–æ–‡å­—
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        merged_doc.add_paragraph(text)
                    else:
                        # å¦‚æœæå–å‡ºæ¥æ˜¯ç©ºçš„ï¼Œå¯èƒ½æ˜¯çº¯å›¾ç‰‡æ‰«æä»¶
                        merged_doc.add_paragraph("[âš ï¸ æç¤ºï¼šæ­¤é¡µå¯èƒ½æ˜¯çº¯å›¾ç‰‡æ‰«æä»¶ï¼Œæœªæå–åˆ°æ–‡å­—]")
                        
        except Exception as e:
            merged_doc.add_paragraph(f"[è¯»å–å¤±è´¥: {str(e)}]")
            st.warning(f"æ–‡ä»¶ {file.name} è¯»å–å‡ºç°é—®é¢˜: {e}")

        # æ’å…¥åˆ†é¡µç¬¦ï¼Œè®©ä¸åŒæ–‡ä»¶çš„å†…å®¹åœ¨è§†è§‰ä¸Šéš”å¼€
        merged_doc.add_page_break()
        progress_bar.progress((i + 1) / total_files)
        
    status_text.text("âœ… æ–‡æ¡£åˆå¹¶å®Œæˆï¼")
    return merged_doc

# ==========================================
# ğŸ”§ å¼•æ“ 2: è¡¨æ ¼åˆå¹¶ (Excels -> Sheets)
# ==========================================
def merge_to_excel_sheets(files):
    output = BytesIO()
    
    # ä½¿ç”¨ xlsxwriter å¼•æ“
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        total_files = len(files)
        
        # è®°å½•ç”¨è¿‡çš„ sheet åï¼Œé˜²æ­¢é‡å¤
        used_sheet_names = {}
        
        for i, file in enumerate(files):
            status_text.text(f"æ­£åœ¨æ¬è¿ ({i+1}/{total_files}): {file.name}")
            
            try:
                # è¯»å– Excel
                df = pd.read_excel(file)
                
                # --- Sheet åå¤„ç†é€»è¾‘ ---
                # Excel sheetåæœ€é•¿31å­—ç¬¦
                base_name = file.name.split('.')[0][:25] 
                # å»é™¤éæ³•å­—ç¬¦
                invalid_chars = ['[', ']', ':', '*', '?', '/', '\\']
                for char in invalid_chars:
                    base_name = base_name.replace(char, '_')
                
                # å¤„ç†é‡å (å¦‚ï¼šåå•.xlsx, åå•.xlsx -> åå•, åå•_1)
                if base_name in used_sheet_names:
                    used_sheet_names[base_name] += 1
                    sheet_name = f"{base_name}_{used_sheet_names[base_name]}"
                else:
                    used_sheet_names[base_name] = 0
                    sheet_name = base_name
                
                # å†™å…¥æ–°çš„ Sheet
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                
            except Exception as e:
                st.error(f"è¡¨æ ¼ {file.name} å¤„ç†å¤±è´¥: {e}")
            
            progress_bar.progress((i + 1) / total_files)
            
    status_text.text("âœ… è¡¨æ ¼åˆå¹¶å®Œæˆï¼")
    return output

# ==========================================
# ğŸš€ ä¸»æ§å°å¸ƒå±€
# ==========================================

st.markdown("---")

col1, col2 = st.columns(2)

# --- å·¦ä¾§ï¼šæ–‡æ¡£åŒº ---
with col1:
    st.subheader("ğŸ“„ æ–‡æ¡£å¤„ç†åŒº")
    st.caption("æ”¯æŒ .docx, .pdf (å«åŠ å¯†) | äº§å‡ºï¼šæ±‡æ€»ç‰ˆ Word")
    doc_files = st.file_uploader("æ‹–å…¥æ–‡æ¡£", type=['docx', 'pdf'], accept_multiple_files=True, key="doc_uploader")
    
    if doc_files:
        if st.button("å¼€å§‹åˆå¹¶æ–‡æ¡£", type="primary"):
            final_doc = merge_to_docx(doc_files)
            
            # ä¿å­˜
            doc_buffer = BytesIO()
            final_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.download_button(
                "ğŸ“¥ ä¸‹è½½è¶…çº§æ–‡æ¡£ (.docx)",
                data=doc_buffer,
                file_name="å…¨é™¢èµ„æ–™æ±‡æ€»_AIç‰ˆ.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# --- å³ä¾§ï¼šè¡¨æ ¼åŒº ---
with col2:
    st.subheader("ğŸ“Š è¡¨æ ¼å¤„ç†åŒº")
    st.caption("æ”¯æŒ .xlsx | äº§å‡ºï¼šå¤š Sheet ç‰ˆ Excel")
    excel_files = st.file_uploader("æ‹–å…¥è¡¨æ ¼", type=['xlsx', 'xls'], accept_multiple_files=True, key="xls_uploader")
    
    if excel_files:
        if st.button("å¼€å§‹åˆå¹¶è¡¨æ ¼"):
            final_excel = merge_to_excel_sheets(excel_files)
            
            st.download_button(
                "ğŸ“¥ ä¸‹è½½è¶…çº§è¡¨æ ¼ (.xlsx)",
                data=final_excel.getvalue(),
                file_name="å…¨é™¢è¡¨æ ¼æ±‡æ€»_å¤šSheetç‰ˆ.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )