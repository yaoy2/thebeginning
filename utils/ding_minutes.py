import configparser
import hashlib
import json
import os
import re
import sqlite3
from dataclasses import dataclass
from datetime import datetime, time, timedelta
from pathlib import Path

import requests
from docx import Document


ROOT_DIR = os.path.dirname(os.path.dirname(__file__))
CONFIG_PATH = os.path.join(ROOT_DIR, "config", "ding_minutes.ini")
DB_PATH = os.path.join(ROOT_DIR, "data", "ding_minutes.db")
CLOUD_EXPORT_PATH = os.path.join(ROOT_DIR, "data", "ding_minutes_cloud.json")

DEFAULT_CONFIG = {
    "watch_dir": r"E:\GoogleDrive\Ding2026",
    "daily_run_time": "19:00",
    "model": "deepseek-v4-pro",
    "api_base": "https://api.deepseek.com",
    "timeout_seconds": "120",
}


def _clean_secret(value):
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def _read_mapping_value(mapping, key, *, clean=True):
    if mapping is None:
        return None
    try:
        value = mapping.get(key)
    except Exception:
        try:
            value = mapping[key]
        except Exception:
            return None
    if clean:
        return _clean_secret(value)
    return value


def get_deepseek_api_key(secrets=None, environ=None):
    environ = os.environ if environ is None else environ

    for key in ("DEEPSEEK_API_KEY", "deepseek_api_key"):
        api_key = _read_mapping_value(secrets, key)
        if api_key:
            return api_key

    deepseek_section = _read_mapping_value(secrets, "deepseek", clean=False)
    if deepseek_section is not None:
        api_key = _read_mapping_value(deepseek_section, "api_key")
        if api_key:
            return api_key

    return _read_mapping_value(environ, "DEEPSEEK_API_KEY")


@dataclass(frozen=True)
class FileStat:
    created_at: datetime
    modified_at: datetime
    size: int


def load_config(path=CONFIG_PATH):
    config = DEFAULT_CONFIG.copy()
    parser = configparser.ConfigParser()
    parser.read(path, encoding="utf-8")
    if parser.has_section("ding_minutes"):
        for key in config:
            if parser.has_option("ding_minutes", key):
                config[key] = parser.get("ding_minutes", key).strip()
    return config


def matches_ding_docx(file_name):
    name = Path(str(file_name)).name
    lower = name.lower()
    if lower.startswith("~$") or not lower.endswith(".docx"):
        return False
    return bool(re.fullmatch(r"export_.+\.docx", lower) or re.fullmatch(r"dt.*\.docx", lower))


def build_scan_window(now=None, run_time_text="19:00"):
    now = now or datetime.now()
    hour_text, minute_text = str(run_time_text).split(":", 1)
    run_time = time(int(hour_text), int(minute_text))
    today_cutoff = datetime.combine(now.date(), run_time)
    end = today_cutoff if now >= today_cutoff else today_cutoff - timedelta(days=1)
    return end - timedelta(days=1), end


def get_file_stat(path):
    path = Path(path)
    stat = path.stat()
    return FileStat(
        created_at=datetime.fromtimestamp(os.path.getctime(path)),
        modified_at=datetime.fromtimestamp(stat.st_mtime),
        size=stat.st_size,
    )


def init_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = get_connection()
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS ding_minutes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            created_at_file TEXT NOT NULL,
            modified_at_file TEXT NOT NULL,
            file_size INTEGER NOT NULL,
            content_hash TEXT NOT NULL,
            original_text TEXT NOT NULL DEFAULT '',
            ai_summary TEXT NOT NULL DEFAULT '',
            remark TEXT NOT NULL DEFAULT '',
            model TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'pending',
            error_message TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            UNIQUE(content_hash)
        )
        """
    )
    conn.commit()
    conn.close()


def get_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA journal_mode=MEMORY")
    conn.row_factory = sqlite3.Row
    return conn


def compute_content_hash(file_path, original_text="", file_stat=None):
    path = Path(file_path)
    stat = file_stat or get_file_stat(path)
    digest = hashlib.sha256()
    digest.update(path.name.lower().encode("utf-8", errors="ignore"))
    digest.update(str(stat.created_at.isoformat()).encode("utf-8"))
    digest.update(str(stat.size).encode("utf-8"))
    digest.update((original_text or "").encode("utf-8", errors="ignore"))
    return digest.hexdigest()


def upsert_file_record(record):
    init_db()
    now = datetime.now().isoformat(timespec="seconds")
    values = {
        "file_name": record.get("file_name", ""),
        "file_path": record.get("file_path", ""),
        "created_at_file": record.get("created_at_file", ""),
        "modified_at_file": record.get("modified_at_file", ""),
        "file_size": int(record.get("file_size", 0) or 0),
        "content_hash": record.get("content_hash", ""),
        "original_text": record.get("original_text", ""),
        "ai_summary": record.get("ai_summary", ""),
        "remark": record.get("remark", ""),
        "model": record.get("model", ""),
        "status": record.get("status", "pending"),
        "error_message": record.get("error_message", ""),
        "created_at": now,
        "updated_at": now,
    }
    conn = get_connection()
    existing = conn.execute(
        "SELECT id FROM ding_minutes WHERE content_hash = ?",
        (values["content_hash"],),
    ).fetchone()
    if existing:
        conn.close()
        return existing["id"]
    cursor = conn.execute(
        """
        INSERT INTO ding_minutes
            (file_name, file_path, created_at_file, modified_at_file, file_size,
             content_hash, original_text, ai_summary, remark, model, status,
             error_message, created_at, updated_at)
        VALUES
            (:file_name, :file_path, :created_at_file, :modified_at_file, :file_size,
             :content_hash, :original_text, :ai_summary, :remark, :model, :status,
             :error_message, :created_at, :updated_at)
        """,
        values,
    )
    conn.commit()
    record_id = cursor.lastrowid
    conn.close()
    return record_id


def update_record(record_id, **fields):
    if not fields:
        return
    allowed = {"ai_summary", "remark", "model", "status", "error_message", "updated_at"}
    safe_fields = {key: value for key, value in fields.items() if key in allowed}
    safe_fields["updated_at"] = datetime.now().isoformat(timespec="seconds")
    assignments = ", ".join(f"{key} = :{key}" for key in safe_fields)
    safe_fields["id"] = record_id
    conn = get_connection()
    conn.execute(f"UPDATE ding_minutes SET {assignments} WHERE id = :id", safe_fields)
    conn.commit()
    conn.close()


def update_remark(record_id, remark):
    update_record(record_id, remark=str(remark or ""))
    sync_cloud_export()


def mark_done(record_id, ai_summary, model):
    update_record(
        record_id,
        ai_summary=ai_summary or "",
        model=model or "",
        status="done",
        error_message="",
    )


def mark_failed(record_id, error_message, status="failed"):
    update_record(record_id, status=status, error_message=_safe_error(error_message))


def get_records(status=None, keyword=None, limit=200):
    init_db()
    where = []
    params = {}
    if status and status != "全部":
        where.append("status = :status")
        params["status"] = status
    if keyword:
        where.append(
            "(file_name LIKE :keyword OR original_text LIKE :keyword OR ai_summary LIKE :keyword OR remark LIKE :keyword)"
        )
        params["keyword"] = f"%{keyword}%"
    where_sql = "WHERE " + " AND ".join(where) if where else ""
    params["limit"] = int(limit)
    conn = get_connection()
    rows = conn.execute(
        f"""
        SELECT * FROM ding_minutes
        {where_sql}
        ORDER BY created_at_file DESC, id DESC
        LIMIT :limit
        """,
        params,
    ).fetchall()
    conn.close()
    return [dict(row) for row in rows]


def get_record(record_id):
    init_db()
    conn = get_connection()
    row = conn.execute("SELECT * FROM ding_minutes WHERE id = ?", (record_id,)).fetchone()
    conn.close()
    return dict(row) if row else None


def get_status_counts():
    init_db()
    conn = get_connection()
    rows = conn.execute("SELECT status, COUNT(*) AS total FROM ding_minutes GROUP BY status").fetchall()
    conn.close()
    return {row["status"]: row["total"] for row in rows}


def build_cloud_payload(records=None):
    records = get_records(limit=1000) if records is None else list(records)
    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "record_count": len(records),
        "records": records,
    }


def sync_cloud_export(path=CLOUD_EXPORT_PATH):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    payload = build_cloud_payload()
    with open(path, "w", encoding="utf-8", newline="\n") as export_file:
        json.dump(payload, export_file, ensure_ascii=False, indent=2)
        export_file.write("\n")
    return path


def load_cloud_export(path=CLOUD_EXPORT_PATH):
    if not os.path.exists(path):
        return {"generated_at": "", "record_count": 0, "records": []}
    with open(path, "r", encoding="utf-8") as export_file:
        payload = json.load(export_file)
    payload["records"] = list(payload.get("records") or [])
    payload["record_count"] = len(payload["records"])
    return payload


def get_cloud_status_counts(records):
    counts = {}
    for record in records:
        status = record.get("status") or "pending"
        counts[status] = counts.get(status, 0) + 1
    return counts


def extract_docx_text(file_path):
    document = Document(file_path)
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def build_ai_prompt(original_text):
    trimmed_text = (original_text or "").strip()
    if len(trimmed_text) > 48000:
        trimmed_text = trimmed_text[:48000] + "\n\n[原文较长，以上为前 48000 字。]"
    return (
        "你是一名高校二级学院行政工作人员的材料整理助手。请把下面的钉钉录音转写整理成"
        "有价值、可归档、可复盘、可继续用于写通知/纪要/新闻/汇报材料的整理稿。\n\n"
        "要求：\n"
        "1. 不使用固定模板，不要机械套用会议纪要八栏目。\n"
        "2. 根据内容自然判断场景、主题和小节标题。\n"
        "3. 清理口水话、重复话、无意义语气词，但保留真实意思。\n"
        "4. 不编造原文没有的信息；不确定处明确写“不确定”。\n"
        "5. 尽量整理出对后续工作有用的信息、观点、结论、待跟进事项和可复用表述。\n"
        "6. 如果原文信息很少，直接说明可整理信息有限。\n\n"
        "原始转写如下：\n"
        f"{trimmed_text}"
    )


def generate_summary_for_record(record_id, ai_client=None, model="deepseek-v4-pro", config=None):
    record = get_record(record_id)
    if not record:
        raise ValueError("记录不存在")
    original_text = (record.get("original_text") or "").strip()
    if not original_text:
        mark_failed(record_id, "原文为空，无法生成整理稿")
        return False
    config = DEFAULT_CONFIG | (config or {})
    client = ai_client or DeepSeekClient(
        api_key=config.get("api_key") or get_deepseek_api_key(),
        model=model or config.get("model", "deepseek-v4-pro"),
        api_base=config.get("api_base", "https://api.deepseek.com"),
        timeout=config.get("timeout_seconds", "120"),
    )
    try:
        summary = client.summarize(original_text)
        mark_done(record_id, summary, model or config.get("model", "deepseek-v4-pro"))
        sync_cloud_export()
        return True
    except Exception as exc:
        mark_failed(record_id, exc)
        return False


class DeepSeekClient:
    def __init__(self, api_key=None, model="deepseek-v4-pro", api_base="https://api.deepseek.com", timeout=120):
        self.api_key = api_key or get_deepseek_api_key()
        self.model = model
        self.api_base = (api_base or "https://api.deepseek.com").rstrip("/")
        self.timeout = int(timeout or 120)

    def summarize(self, text):
        if not self.api_key:
            raise RuntimeError("未配置 DEEPSEEK_API_KEY")
        response = requests.post(
            f"{self.api_base}/chat/completions",
            headers={
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
            },
            data=json.dumps(
                {
                    "model": self.model,
                    "messages": [
                        {"role": "system", "content": "你擅长整理中文行政工作录音转写。"},
                        {"role": "user", "content": build_ai_prompt(text)},
                    ],
                    "temperature": 0.3,
                },
                ensure_ascii=False,
            ).encode("utf-8"),
            timeout=self.timeout,
        )
        if response.status_code >= 400:
            raise RuntimeError(f"DeepSeek 请求失败：HTTP {response.status_code}")
        payload = response.json()
        try:
            return payload["choices"][0]["message"]["content"].strip()
        except (KeyError, IndexError, TypeError) as exc:
            raise RuntimeError("DeepSeek 返回格式异常") from exc


def scan_once(config=None, now=None, ai_client=None, stat_func=get_file_stat):
    config = DEFAULT_CONFIG | (config or load_config())
    init_db()
    start, end = build_scan_window(now=now, run_time_text=config.get("daily_run_time", "19:00"))
    watch_dir = Path(config.get("watch_dir", "")).expanduser()
    result = {"found": 0, "processed": 0, "skipped": 0, "failed": 0, "window_start": start, "window_end": end}
    if not watch_dir.exists():
        result["error"] = f"文件夹不存在：{watch_dir}"
        return result

    client = ai_client
    api_key = config.get("api_key") or get_deepseek_api_key()
    if client is None and api_key:
        client = DeepSeekClient(
            api_key=api_key,
            model=config.get("model", "deepseek-v4-pro"),
            api_base=config.get("api_base", "https://api.deepseek.com"),
            timeout=config.get("timeout_seconds", "120"),
        )

    for file_path in sorted(watch_dir.glob("*.docx")):
        if not matches_ding_docx(file_path.name):
            result["skipped"] += 1
            continue
        try:
            file_stat = stat_func(file_path)
            if not (start <= file_stat.created_at <= end):
                result["skipped"] += 1
                continue
            result["found"] += 1
            original_text = extract_docx_text(file_path)
            content_hash = compute_content_hash(file_path, original_text, file_stat)
            record_id = upsert_file_record(
                {
                    "file_name": file_path.name,
                    "file_path": str(file_path),
                    "created_at_file": file_stat.created_at.isoformat(timespec="seconds"),
                    "modified_at_file": file_stat.modified_at.isoformat(timespec="seconds"),
                    "file_size": file_stat.size,
                    "content_hash": content_hash,
                    "original_text": original_text,
                    "model": config.get("model", ""),
                    "status": "pending",
                }
            )
            existing = next((record for record in get_records(limit=500) if record["id"] == record_id), None)
            if existing and existing.get("status") == "done":
                result["skipped"] += 1
                continue
            if not original_text:
                mark_failed(record_id, "Word 正文为空或无法提取")
                result["failed"] += 1
                continue
            if client is None:
                mark_failed(record_id, "未配置 DEEPSEEK_API_KEY，已登记原文，待生成整理稿", status="pending")
                result["processed"] += 1
                continue
            ai_summary = client.summarize(original_text)
            mark_done(record_id, ai_summary, config.get("model", "deepseek-v4-pro"))
            result["processed"] += 1
        except Exception as exc:
            result["failed"] += 1
            try:
                mark_failed(record_id, exc)
            except Exception:
                pass
    sync_cloud_export()
    return result


def _safe_error(error):
    text = str(error or "")
    text = re.sub(r"Bearer\s+[A-Za-z0-9._-]+", "Bearer [hidden]", text)
    text = re.sub(r"sk-[A-Za-z0-9_-]+", "sk-[hidden]", text)
    return text[:500]
